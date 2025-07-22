import streamlit as st
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml import parse_xml
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib import colors
import io
from datetime import datetime
import re

# Page configuration
st.set_page_config(
    page_title="Multi-Template Resume Generator",
    page_icon="📄",
    layout="wide"
)

st.title("📄 Multi-Template Resume Generator")
st.markdown("Choose from 5 professionally designed resume templates with centered contact information!")

# Resume Template Definitions
RESUME_TEMPLATES = {
    "Classic Professional": {
        "description": "Traditional black and white with clean lines - Universally accepted",
        "color_scheme": {
            "primary": RGBColor(0, 0, 0),  # Black
            "secondary": RGBColor(64, 64, 64),  # Dark Gray
            "accent": RGBColor(128, 128, 128)  # Light Gray
        },
        "pdf_colors": {
            "primary": colors.black,
            "secondary": colors.Color(0.25, 0.25, 0.25),
            "accent": colors.Color(0.5, 0.5, 0.5)
        },
        "font_style": "Arial",
        "header_style": "underlined"
    },
    "Modern Blue": {
        "description": "Contemporary design with professional blue accents - Tech-friendly",
        "color_scheme": {
            "primary": RGBColor(0, 51, 102),  # Navy Blue
            "secondary": RGBColor(0, 102, 204),  # Medium Blue
            "accent": RGBColor(102, 153, 255)  # Light Blue
        },
        "pdf_colors": {
            "primary": colors.Color(0, 0.2, 0.4),
            "secondary": colors.Color(0, 0.4, 0.8),
            "accent": colors.Color(0.4, 0.6, 1)
        },
        "font_style": "Calibri",
        "header_style": "colored_background"
    },
    "Executive Green": {
        "description": "Sophisticated green theme for senior positions - Leadership-focused",
        "color_scheme": {
            "primary": RGBColor(0, 100, 0),  # Dark Green
            "secondary": RGBColor(34, 139, 34),  # Forest Green
            "accent": RGBColor(144, 238, 144)  # Light Green
        },
        "pdf_colors": {
            "primary": colors.Color(0, 0.4, 0),
            "secondary": colors.Color(0.13, 0.55, 0.13),
            "accent": colors.Color(0.56, 0.93, 0.56)
        },
        "font_style": "Times New Roman",
        "header_style": "bold_colored"
    },
    "Creative Purple": {
        "description": "Stylish purple design for creative professionals - Artistic appeal",
        "color_scheme": {
            "primary": RGBColor(75, 0, 130),  # Indigo
            "secondary": RGBColor(138, 43, 226),  # Blue Violet
            "accent": RGBColor(221, 160, 221)  # Plum
        },
        "pdf_colors": {
            "primary": colors.Color(0.29, 0, 0.51),
            "secondary": colors.Color(0.54, 0.17, 0.89),
            "accent": colors.Color(0.87, 0.63, 0.87)
        },
        "font_style": "Georgia",
        "header_style": "gradient_effect"
    },
    "Warm Orange": {
        "description": "Energetic orange theme for dynamic professionals - Marketing-friendly",
        "color_scheme": {
            "primary": RGBColor(204, 85, 0),  # Dark Orange
            "secondary": RGBColor(255, 140, 0),  # Dark Orange
            "accent": RGBColor(255, 218, 185)  # Peach
        },
        "pdf_colors": {
            "primary": colors.Color(0.8, 0.33, 0),
            "secondary": colors.Color(1, 0.55, 0),
            "accent": colors.Color(1, 0.85, 0.73)
        },
        "font_style": "Verdana",
        "header_style": "boxed"
    }
}

def rgbcolor_to_rgb(rgbcolor):
    """Convert docx RGBColor to (r,g,b) tuple"""
    color_hex = str(rgbcolor)  # Gets hex string like "003366"
    r = int(color_hex[0:2], 16)
    g = int(color_hex[2:4], 16) 
    b = int(color_hex[4:6], 16)
    return (r, g, b)

def format_url(url):
    """Format URL to ensure it has proper protocol"""
    if not url or url.strip() == "":
        return ""
    
    url = url.strip()
    if not url.startswith(('http://', 'https://')):
        if 'linkedin.com' in url or 'github.com' in url:
            url = 'https://' + url
        else:
            url = 'https://' + url
    return url

def add_hyperlink_to_paragraph(paragraph, text, url):
    """Add a hyperlink to a paragraph in Word document"""
    # Create hyperlink
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create a new run object and add it to the hyperlink
    new_run = OxmlElement('w:r')
    
    # Create run properties for styling
    rPr = OxmlElement('w:rPr')
    
    # Add color (blue) and underline for hyperlink styling
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0066CC')  # Blue color
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    
    # Create text element
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink

def add_colored_line_after_paragraph(paragraph, color_rgb):
    """Add a colored horizontal line after a paragraph in Word"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr, 'w:shd', 'w:tabs', 'w:suppressAutoHyphens')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    # Convert RGBColor to hex
    color_hex = f'{rgbcolor_to_rgb(color_rgb)[0]:02x}{rgbcolor_to_rgb(color_rgb)[1]:02x}{rgbcolor_to_rgb(color_rgb)[2]:02x}'
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)

def create_template_word_doc(data, template_name):
    """Create a Word document with template-specific styling and CENTERED contact info"""
    template_config = RESUME_TEMPLATES[template_name]
    colors = template_config["color_scheme"]
    font_name = template_config["font_style"]
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(10)
    
    # NAME SECTION - ALWAYS CENTERED
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(data['name'].upper())
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.name = font_name
    name_run.font.color.rgb = colors["primary"]
    name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # ALWAYS CENTER
    name_para.space_after = Pt(6)
    
    # CONTACT INFO - ALWAYS CENTERED with clickable links
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # ALWAYS CENTER
    
    # Add email
    email_run = contact_para.add_run(f"📧 {data['email']}")
    email_run.font.size = Pt(9)
    email_run.font.color.rgb = colors["secondary"]
    
    # Add phone
    phone_run = contact_para.add_run(f"    📞 {data['phone']}")
    phone_run.font.size = Pt(9)
    phone_run.font.color.rgb = colors["secondary"]
    
    # Add location
    location_run = contact_para.add_run(f"    📍 {data['location']}")
    location_run.font.size = Pt(9)
    location_run.font.color.rgb = colors["secondary"]
    
    # Add LinkedIn as clickable link
    if data.get('linkedin') and data['linkedin'].strip():
        linkedin_text_run = contact_para.add_run("    🔗 ")
        linkedin_text_run.font.size = Pt(9)
        linkedin_text_run.font.color.rgb = colors["secondary"]
        
        # Add clickable "LinkedIn" text
        linkedin_url = format_url(data['linkedin'])
        add_hyperlink_to_paragraph(contact_para, "LinkedIn", linkedin_url)
    
    # Add GitHub as clickable link
    if data.get('github') and data['github'].strip():
        github_text_run = contact_para.add_run("    💻 ")
        github_text_run.font.size = Pt(9)
        github_text_run.font.color.rgb = colors["secondary"]
        
        # Add clickable "GitHub" text
        github_url = format_url(data['github'])
        add_hyperlink_to_paragraph(contact_para, "GitHub", github_url)
    
    contact_para.space_after = Pt(12)
    
    # Add styled line after header
    add_colored_line_after_paragraph(contact_para, colors["primary"])
    
    # Helper function to create section headers
    def create_section_header(title):
        heading = doc.add_paragraph()
        heading.space_before = Pt(16)
        heading_run = heading.add_run(title.upper())
        heading_run.bold = True
        heading_run.font.size = Pt(12)
        heading_run.font.name = font_name
        heading_run.font.color.rgb = colors["primary"]
        
        # Template-specific styling
        if template_config["header_style"] == "underlined":
            heading_run.underline = True
        elif template_config["header_style"] == "colored_background":
            # Add background shading (simulated with border)
            add_colored_line_after_paragraph(heading, colors["accent"])
        
        heading.space_after = Pt(8)
        return heading
    
    # EDUCATION SECTION
    create_section_header("Education")
    edu_lines = data['education'].strip().split('\n')
    for line in edu_lines:
        if line.strip():
            edu_para = doc.add_paragraph()
            if any(word in line.lower() for word in ['university', 'college', 'school', 'institute']):
                edu_run = edu_para.add_run(line.strip())
                edu_run.bold = True
                edu_run.font.size = Pt(10)
                edu_run.font.color.rgb = colors["secondary"]
            else:
                edu_run = edu_para.add_run(line.strip())
                edu_run.font.size = Pt(10)
            edu_para.space_after = Pt(3)
    
    # PROJECTS SECTION
    if data.get('projects') and data['projects'].strip():
        create_section_header("Projects")
        proj_lines = data['projects'].strip().split('\n')
        
        for line in proj_lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('•') or line.startswith('-'):
                bullet_para = doc.add_paragraph()
                bullet_para.style = 'List Bullet'
                bullet_run = bullet_para.add_run(line[1:].strip())
                bullet_run.font.size = Pt(10)
                bullet_para.space_after = Pt(2)
            else:
                proj_para = doc.add_paragraph()
                proj_run = proj_para.add_run(line)
                proj_run.bold = True
                proj_run.font.size = Pt(10)
                proj_run.font.color.rgb = colors["secondary"]
                proj_para.space_after = Pt(3)
    
    # EXPERIENCE SECTION
    create_section_header("Professional Experience")
    if data.get('experience'):
        exp_lines = data['experience'].strip().split('\n')
        for line in exp_lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('•') or line.startswith('-'):
                bullet_para = doc.add_paragraph()
                bullet_para.style = 'List Bullet'
                bullet_run = bullet_para.add_run(line[1:].strip())
                bullet_run.font.size = Pt(10)
                bullet_para.space_after = Pt(2)
            else:
                job_para = doc.add_paragraph()
                job_run = job_para.add_run(line)
                job_run.bold = True
                job_run.font.size = Pt(10)
                job_run.font.color.rgb = colors["secondary"]
                job_para.space_after = Pt(3)
    
    # ACHIEVEMENTS SECTION
    if data.get('achievements') and data['achievements'].strip():
        create_section_header("Achievements")
        ach_lines = data['achievements'].strip().split('\n')
        for line in ach_lines:
            if line.strip():
                ach_para = doc.add_paragraph()
                ach_run = ach_para.add_run(f"• {line.strip()}")
                ach_run.font.size = Pt(10)
                ach_para.space_after = Pt(3)
    
    # TECHNICAL SKILLS SECTION
    if data.get('skills') and data['skills'].strip():
        create_section_header("Technical Skills")
        skills_lines = data['skills'].strip().split('\n')
        for line in skills_lines:
            if line.strip():
                skills_para = doc.add_paragraph()
                if ':' in line:
                    parts = line.split(':', 1)
                    category_run = skills_para.add_run(parts[0].strip() + ': ')
                    category_run.bold = True
                    category_run.font.size = Pt(10)
                    category_run.font.color.rgb = colors["secondary"]
                    skills_run = skills_para.add_run(parts[1].strip())
                    skills_run.font.size = Pt(10)
                else:
                    skills_run = skills_para.add_run(line.strip())
                    skills_run.font.size = Pt(10)
                skills_para.space_after = Pt(3)
    
    # Save to BytesIO
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def create_template_pdf(data, template_name):
    """Create a PDF with template-specific styling and CENTERED contact info"""
    template_config = RESUME_TEMPLATES[template_name]
    pdf_colors = template_config["pdf_colors"]
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=letter,
        rightMargin=0.75*inch,
        leftMargin=0.75*inch,
        topMargin=0.5*inch,
        bottomMargin=0.5*inch
    )
    
    styles = getSampleStyleSheet()
    story = []
    
    # Custom styles with template colors - ALWAYS CENTERED FOR CONTACT INFO
    name_style = ParagraphStyle(
        'NameStyle',
        parent=styles['Normal'],
        fontSize=18,
        spaceAfter=6,
        spaceBefore=0,
        textColor=pdf_colors["primary"],
        fontName='Helvetica-Bold',
        alignment=1  # ALWAYS CENTER (1 = center alignment)
    )
    
    contact_style = ParagraphStyle(
        'ContactStyle',
        parent=styles['Normal'],
        fontSize=9,
        spaceAfter=12,
        spaceBefore=2,
        textColor=pdf_colors["secondary"],
        alignment=1  # ALWAYS CENTER (1 = center alignment)
    )
    
    section_style = ParagraphStyle(
        'SectionStyle',
        parent=styles['Normal'],
        fontSize=12,
        spaceAfter=8,
        spaceBefore=16,
        textColor=pdf_colors["primary"],
        fontName='Helvetica-Bold'
    )
    
    content_style = ParagraphStyle(
        'ContentStyle',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=3,
        spaceBefore=0
    )
    
    job_style = ParagraphStyle(
        'JobStyle',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=3,
        spaceBefore=0,
        textColor=pdf_colors["secondary"],
        fontName='Helvetica-Bold'
    )
    
    # NAME - ALWAYS CENTERED
    story.append(Paragraph(f'<b>{data["name"].upper()}</b>', name_style))
    
    # CONTACT INFO - ALWAYS CENTERED with clickable links
    contact_parts = [f'{data["email"]}', f'{data["phone"]}', f'{data["location"]}']
    
    # Add LinkedIn as clickable link
    if data.get('linkedin') and data['linkedin'].strip():
        linkedin_url = format_url(data['linkedin'])
        contact_parts.append(f'<link href="{linkedin_url}" color="blue">LinkedIn</link>')
    
    # Add GitHub as clickable link  
    if data.get('github') and data['github'].strip():
        github_url = format_url(data['github'])
        contact_parts.append(f'<link href="{github_url}" color="blue">GitHub</link>')
    
    contact_info = ' | '.join(contact_parts)
    story.append(Paragraph(contact_info, contact_style))
    
    # Add separator line
    story.append(Spacer(1, 6))
    
    # Helper function to add colored section headers
    def add_section_header(title):
        if template_config["header_style"] == "colored_background":
            # Create a table for background effect
            header_table = Table([[title.upper()]], colWidths=[7*inch])
            header_table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,-1), pdf_colors["accent"]),
                ('TEXTCOLOR', (0,0), (-1,-1), pdf_colors["primary"]),
                ('FONTNAME', (0,0), (-1,-1), 'Helvetica-Bold'),
                ('FONTSIZE', (0,0), (-1,-1), 12),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('LEFTPADDING', (0,0), (-1,-1), 6),
                ('RIGHTPADDING', (0,0), (-1,-1), 6),
                ('TOPPADDING', (0,0), (-1,-1), 4),
                ('BOTTOMPADDING', (0,0), (-1,-1), 4),
            ]))
            story.append(header_table)
            story.append(Spacer(1, 8))
        else:
            story.append(Paragraph(f'<b>{title.upper()}</b>', section_style))
    
    # EDUCATION
    add_section_header("Education")
    edu_lines = data['education'].strip().split('\n')
    for line in edu_lines:
        if line.strip():
            if any(word in line.lower() for word in ['university', 'college', 'school', 'institute']):
                story.append(Paragraph(f'<b>{line.strip()}</b>', job_style))
            else:
                story.append(Paragraph(line.strip(), content_style))
    
    # PROJECTS
    if data.get('projects') and data['projects'].strip():
        add_section_header("Projects")
        proj_lines = data['projects'].strip().split('\n')
        for line in proj_lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith('•') or line.startswith('-'):
                story.append(Paragraph(f'• {line[1:].strip()}', content_style))
            else:
                story.append(Paragraph(f'<b>{line}</b>', job_style))
    
    # EXPERIENCE
    add_section_header("Professional Experience")
    if data.get('experience'):
        exp_lines = data['experience'].strip().split('\n')
        for line in exp_lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith('•') or line.startswith('-'):
                story.append(Paragraph(f'• {line[1:].strip()}', content_style))
            else:
                story.append(Paragraph(f'<b>{line}</b>', job_style))
    
    # ACHIEVEMENTS
    if data.get('achievements') and data['achievements'].strip():
        add_section_header("Achievements")
        ach_lines = data['achievements'].strip().split('\n')
        for line in ach_lines:
            if line.strip():
                story.append(Paragraph(f'• {line.strip()}', content_style))
    
    # TECHNICAL SKILLS
    if data.get('skills') and data['skills'].strip():
        add_section_header("Technical Skills")
        skills_lines = data['skills'].strip().split('\n')
        for line in skills_lines:
            if line.strip():
                if ':' in line:
                    parts = line.split(':', 1)
                    # Create RGB values for PDF colors
                    sec_rgb = pdf_colors["secondary"]
                    color_r = int(sec_rgb.red * 255)
                    color_g = int(sec_rgb.green * 255)
                    color_b = int(sec_rgb.blue * 255)
                    skill_para = Paragraph(f'<b><font color="rgb({color_r},{color_g},{color_b})">{parts[0].strip()}:</font></b> {parts[1].strip()}', content_style)
                    story.append(skill_para)
                else:
                    story.append(Paragraph(line.strip(), content_style))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

# Initialize session state
if 'resume_data' not in st.session_state:
    st.session_state.resume_data = None
if 'selected_template' not in st.session_state:
    st.session_state.selected_template = "Classic Professional"

# TEMPLATE SELECTION SIDEBAR
st.sidebar.header("🎨 Choose Your Resume Template")

selected_template = st.sidebar.selectbox(
    "Select Template Style",
    list(RESUME_TEMPLATES.keys()),
    index=list(RESUME_TEMPLATES.keys()).index(st.session_state.selected_template)
)

st.session_state.selected_template = selected_template

# Display template info
template_info = RESUME_TEMPLATES[selected_template]
st.sidebar.markdown(f"**{selected_template}**")
st.sidebar.markdown(template_info["description"])
st.sidebar.markdown(f"**Font:** {template_info['font_style']}")
st.sidebar.markdown(f"**Style:** {template_info['header_style'].replace('_', ' ').title()}")

# Template preview colors with fixed RGB conversion
colors_scheme = template_info["color_scheme"]
st.sidebar.markdown("**Color Scheme:**")

# Convert RGBColors to RGB tuples
primary_rgb = rgbcolor_to_rgb(colors_scheme["primary"])
secondary_rgb = rgbcolor_to_rgb(colors_scheme["secondary"])
accent_rgb = rgbcolor_to_rgb(colors_scheme["accent"])

col1, col2, col3 = st.sidebar.columns(3)
with col1:
    st.markdown(f'<div style="background-color:rgb({primary_rgb[0]},{primary_rgb[1]},{primary_rgb[2]}); height:20px; border-radius:3px;"></div>', unsafe_allow_html=True)
    st.caption("Primary")
with col2:
    st.markdown(f'<div style="background-color:rgb({secondary_rgb[0]},{secondary_rgb[1]},{secondary_rgb[2]}); height:20px; border-radius:3px;"></div>', unsafe_allow_html=True)
    st.caption("Secondary")
with col3:
    st.markdown(f'<div style="background-color:rgb({accent_rgb[0]},{accent_rgb[1]},{accent_rgb[2]}); height:20px; border-radius:3px;"></div>', unsafe_allow_html=True)
    st.caption("Accent")

# Highlight the centered alignment feature
st.sidebar.markdown("---")
st.sidebar.info("📍 **Contact Info**: Name, email, phone, location, LinkedIn, and GitHub are **always centered** in all templates for professional appearance!")

# MAIN INTERFACE
col1, col2 = st.columns([1, 1])

with col1:
    st.header(f"📝 {selected_template} Resume")
    
    # Personal Information
    st.subheader("👤 Personal Information")
    st.info("📍 All contact information will be **centered** in your resume")
    
    name = st.text_input("Full Name *", placeholder="Your Full Name")
    email = st.text_input("Email *", placeholder="your.email@example.com")
    phone = st.text_input("Phone *", placeholder="+1 (555) 123-4567")
    location = st.text_input("Location *", placeholder="City, State")
    
    # Enhanced LinkedIn and GitHub inputs with better descriptions
    st.markdown("**🔗 Professional Links** (will appear as clickable words)")
    linkedin = st.text_input(
        "LinkedIn Profile", 
        placeholder="linkedin.com/in/yourname or full URL",
        help="Enter your LinkedIn URL (with or without https://). It will appear as 'LinkedIn' in your resume."
    )
    github = st.text_input(
        "GitHub Profile", 
        placeholder="github.com/yourname or full URL",
        help="Enter your GitHub URL (with or without https://). It will appear as 'GitHub' in your resume."
    )
    
    # Education
    st.subheader("🎓 Education")
    education = st.text_area(
        "Education Details *", 
        placeholder="""University Name
Bachelor/Master of [Degree] in [Field]
GPA: X.XX/4.0 (if above 3.5)
Graduation: Month Year""",
        height=100
    )
    
    # Projects
    st.subheader("🚀 Projects")
    projects = st.text_area(
        "Projects", 
        placeholder="""Project Name:
• Brief description of what the project does
• Technologies used and your role
• Key achievements or results

Another Project:
• Description with impact and results
• Technical stack and methodologies used""",
        height=150
    )
    
    # Experience
    st.subheader("💼 Professional Experience")
    experience = st.text_area(
        "Experience", 
        placeholder="""Job Title - Company Name (Start Date - End Date)
• Achieved specific result using particular method/technology
• Led/developed/improved something with quantifiable impact
• Collaborated with team on important project or initiative""",
        height=120
    )
    
    # Achievements
    st.subheader("🏆 Achievements")
    achievements = st.text_area(
        "Achievements", 
        placeholder="""Award/Recognition - Description and year
Publication - Title and publication details
Competition - Placement and competition name
Certification - Name and issuing organization""",
        height=100
    )
    
    # Technical Skills
    st.subheader("💻 Technical Skills")
    skills = st.text_area(
        "Technical Skills", 
        placeholder="""Programming Languages: Python, Java, JavaScript
Frameworks & Libraries: React, Django, TensorFlow
Databases: MySQL, PostgreSQL, MongoDB
Tools & Technologies: Git, Docker, AWS""",
        height=100
    )

with col2:
    st.header("📄 Resume Preview & Download")
    
    # Show centered alignment preview
    st.success("📍 **Centered Contact Section**: Your name and contact information will always be centered for a professional appearance!")
    
    # Show link preview if provided
    if linkedin or github:
        st.info("🔗 **Link Preview**: Your links will appear as clickable words in the resume:")
        if linkedin:
            st.markdown(f"- **LinkedIn** → {format_url(linkedin)}")
        if github:
            st.markdown(f"- **GitHub** → {format_url(github)}")
    
    if st.button(f"🚀 Generate {selected_template} Resume", type="primary"):
        # Validation
        required_fields = [name, email, phone, location, education]
        if all(field.strip() for field in required_fields):
            resume_data = {
                'name': name,
                'email': email,
                'phone': phone,
                'location': location,
                'linkedin': linkedin,
                'github': github,
                'education': education,
                'projects': projects,
                'experience': experience,
                'achievements': achievements,
                'skills': skills
            }
            
            st.success(f"✅ {selected_template} resume generated successfully with centered contact info!")
            st.session_state.resume_data = resume_data
            
        else:
            st.error("❌ Please fill all required fields (marked with *)")
    
    # Download section
    if st.session_state.resume_data:
        st.subheader("📥 Download Your Resume")
        
        col_word, col_pdf = st.columns(2)
        
        with col_word:
            word_doc = create_template_word_doc(st.session_state.resume_data, selected_template)
            # Fixed file name generation
            clean_name = st.session_state.resume_data['name'].replace(' ', '_')
            clean_template = selected_template.replace(' ', '_')
            word_filename = f"{clean_name}_{clean_template}_Resume.docx"
            
            st.download_button(
                label="📄 Download Word (.docx)",
                data=word_doc.getvalue(),
                file_name=word_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        
        with col_pdf:
            pdf_doc = create_template_pdf(st.session_state.resume_data, selected_template)
            # Fixed file name generation
            clean_name = st.session_state.resume_data['name'].replace(' ', '_')
            clean_template = selected_template.replace(' ', '_')
            pdf_filename = f"{clean_name}_{clean_template}_Resume.pdf"
            
            st.download_button(
                label="📑 Download PDF",
                data=pdf_doc.getvalue(),
                file_name=pdf_filename,
                mime="application/pdf",
                use_container_width=True
            )
        
        # Template comparison
        st.subheader("🔄 Try Other Templates")
        st.markdown("**Click to switch templates and see different styles:**")
        
        template_cols = st.columns(5)
        for i, (template_name, template_info) in enumerate(RESUME_TEMPLATES.items()):
            with template_cols[i]:
                if st.button(template_name.split()[0], key=f"switch_{i}", help=template_info["description"]):
                    st.session_state.selected_template = template_name
                    st.rerun()

# Template Comparison Table
st.markdown("---")
st.subheader("📊 Template Comparison")

comparison_data = []
for template_name, details in RESUME_TEMPLATES.items():
    comparison_data.append({
        "Template": template_name,
        "Best For": details["description"].split(" - ")[1] if " - " in details["description"] else details["description"],
        "Font Style": details["font_style"],
        "Header Style": details["header_style"].replace('_', ' ').title(),
        "Contact Alignment": "Always Centered ✓"
    })

st.table(comparison_data)

# Enhanced Tips Section
st.markdown("---")
st.markdown("### 💡 Template Selection Guide")

tips_col1, tips_col2, tips_col3 = st.columns(3)

with tips_col1:
    st.markdown("""
    **🎯 Industry Recommendations:**
    - **Classic Professional**: Banking, Law, Healthcare
    - **Modern Blue**: Technology, Engineering
    - **Executive Green**: Management, Finance, Consulting
    """)

with tips_col2:
    st.markdown("""
    **🎨 Visual Impact:**
    - **Creative Purple**: Design, Marketing, Media
    - **Warm Orange**: Sales, Marketing, Startups
    - **All templates**: Professional centered header
    """)

with tips_col3:
    st.markdown("""
    **📍 Professional Layout:**
    - Contact info always centered for elegance
    - Clickable LinkedIn and GitHub links
    - ATS-friendly structure maintained
    - Clean, professional appearance
    """)

st.success("🎯 **Perfect!** All contact information (Name, Email, Phone, Location, LinkedIn, GitHub) is now **always centered** in every template for maximum professional impact!")
